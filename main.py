# 自動生成機房紀錄表word檔案程式
# 思路
# 1 版面編排
# 2 插入表格
# 3 篩選日期: 查看date_caculate.py

# -*- coding: UTF8 -*-
from docx import Document
# 字體大小, 欄位大小單位
from docx.shared import Cm, Pt
# 標楷體字型設定
from docx.oxml.ns import qn
# 文字置中
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# 表格文字置中
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
# 計算日期的函數
from date_caculate import getMothDate
def setFont(doc):
# 標楷體字型設定
    doc.styles['Normal'].font.name = u'標楷體'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
def add_logo(paragraph):
    # 加入圖片
    logo_run = paragraph.add_run()
    logo_run.add_picture('kaofeng_ico.jpg',width=Cm(1.1))
def add_title(paragraph):
    run = paragraph.add_run()
    run1 = paragraph.add_run()
    run.font.size = Pt(28)
    run.text = "高鳳國際物流股份有限公司\n"
    run1.font.size = Pt(20)
    run1.text = "2021年機房檢查記錄表"
def fill_in_data(table):
    for i in range(len(need_to_record_date)):
        table.cell(i, 0).text = need_to_record_date[i][0]
        table.cell(i, 2).text = need_to_record_date[i][1]
def set_table_font_and_font_size(table, fontsize):
    text_size = fontsize
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(text_size)
# 設定要生成的年月份
year = 2021
month = 9

need_to_record_date = getMothDate(year, month)
doc = Document()
# 設定字型
setFont(doc)
paragraph = doc.add_paragraph()
# 標題文字置中
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# 加入高鳳Logo
add_logo(paragraph)
# 加入標題文字
add_title(paragraph)
# 版面編排
# 設定欄寬
col_width_dic = {0: 2.52, 1: 2.75, 2: 2.54, 3: 5.72, 4: 2.54, 5: 2.64}# 總寬18.71公分
# 設定欄位標題
col_text_dic = {0: "日期時間", 1: "溫度/濕度", 2: "備份位置", 3: "狀態說明", 4: "登記人員", 5: "單位主管"}
# 加入表格
headtable = doc.add_table(1,6, style='Table Grid')
headtable.alignment = WD_TABLE_ALIGNMENT.CENTER
table = doc.add_table(len(need_to_record_date),6, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 標題欄位高度1.27cm
for row in headtable.rows: row.height=Cm(1.27)
# 其他行高度0.74cm
for row in table.rows: row.height=Cm(0.71)

for col_num in range(6):
    # 根據col_width_dic 設定欄寬
    headtable.cell(0, col_num).width = Cm(col_width_dic[col_num])
    table.cell(0, col_num).width = Cm(col_width_dic[col_num])
    # 根據col_text_dic 設定欄位標題文字
    headtable.cell(0, col_num).text = col_text_dic[col_num]
    # 設定欄位標題文字置中對齊
    headtable.cell(0, col_num).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    headtable.cell(0, col_num).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
# 設定欄位標題文字大小
set_table_font_and_font_size(headtable, 14)
# 填入日期與磁帶編號資料
fill_in_data(table)
# 存檔
doc.save("new.doc")

# 成功打包成exe 的參數設定:
# pip install pyinstaller==3.4
# pip install pyqt5==5.9.2
# 參考網站:
# https://article.itxueyuan.com/kyZej9
# pyinstaller -F main.py